import os
import pymorphy3
from docx import Document
from django.conf import settings
from datetime import datetime


class ContractGenerator:
    """Генератор договоров"""

    def __init__(self):
        self.morph = pymorphy3.MorphAnalyzer()
        self.months_ru = [
            "января", "февраля", "марта", "апреля", "мая", "июня",
            "июля", "августа", "сентября", "октября", "ноября", "декабря"
        ]

    def genitive_case(self, word):
        """Склонение слова в родительный падеж"""
        if not word:
            return ""

        try:
            words = word.split()
            modified_words = []

            for w in words:
                parsed_word = self.morph.parse(w)[0]
                if "NOUN" in parsed_word.tag or "ADJF" in parsed_word.tag:
                    genitive = parsed_word.inflect({"gent"}).word
                    modified_words.append(genitive)
                else:
                    modified_words.append(w)

            if len(words) == 3:
                if 'уляшва' in modified_words:
                    modified_words[0] = 'уляшова'
                capitalized_words = [word.capitalize() for word in modified_words]
                return " ".join(capitalized_words)
            else:
                return " ".join(modified_words)

        except Exception as e:
            print(f"Ошибка при склонении слова '{word}': {e}")
            return word

    def format_full_name(self, surname, name, patronymic):
        """Форматирование ФИО в формате 'Фамилия И.О.'"""
        return f"{surname.capitalize()} {name[0].upper()}. {patronymic[0].upper()}."

    def get_workday_phrase(self, number):
        """Формирование фразы о рабочих днях"""
        # Определение формы слова "рабочий"
        if number % 10 == 1 and number % 100 != 11:
            workday_form = self.morph.parse('рабочий')[0].inflect(
                {'nomn', 'sing'}).word  # Именительный падеж, единственное число
        elif number % 10 in [2, 3, 4] and not (number % 100 in [12, 13, 14]):
            workday_form = self.morph.parse('рабочий')[0].inflect(
                {'gent', 'plur'}).word  # Родительный падеж, множественное число
        else:
            workday_form = self.morph.parse('рабочий')[0].inflect(
                {'gent', 'plur'}).word  # Родительный падеж, множественное число

        # Определение формы слова "день"
        if number % 10 == 1 and number % 100 != 11:
            day_form = 'день'  # 1, 21
        elif number % 10 in [2, 3, 4] and not (number % 100 in [12, 13, 14]):
            day_form = 'дня'  # 2, 3, 4, 22, 23, 24
        else:
            day_form = 'дней'  # 5-20, 25-31

        # Формируем строку
        return f"{number} {workday_form} {day_form}"

    def generate_contract(self, legal_entity, organization, timeframe=21):
        """Генерация договора"""
        now = datetime.now()
        day_of_month = now.day
        month_num = now.month
        month_name = self.months_ru[month_num - 1]

        c_number = f'{day_of_month}/{month_num}/{str(now.year)[2::]}/36/{organization.inn[-4:]}'

        data = {
            'юл': legal_entity.name.upper(),
            'юл_огрн': legal_entity.ogrn,
            'юл_инн': legal_entity.inn,
            'юл_должность': legal_entity.ceo_title,
            'юл_должность_рп': self.genitive_case(legal_entity.ceo_title),
            'юл_фио': legal_entity.ceo_name,
            'юл_фио_рп': self.genitive_case(legal_entity.ceo_name),
            'юл_фио_кратко': self.format_full_name(*legal_entity.ceo_name.split()),
            'юл_кпп': legal_entity.kpp,
            'юл_рс': legal_entity.r_s,
            'юл_банк': legal_entity.bank,
            'юл_бик': legal_entity.bik,
            'юл_корс': legal_entity.k_s,
            'юл_адрес': legal_entity.address,
            'юл_email': legal_entity.email,
            'орг': organization.name.upper(),
            'орг_огрн': organization.ogrn,
            'орг_инн': organization.inn,
            'инн_4': organization.inn[-4:],
            'орг_должность': organization.ceo_title,
            'орг_должность_рп': self.genitive_case(organization.ceo_title),
            'орг_фио': organization.ceo_name,
            'орг_фио_рп': self.genitive_case(organization.ceo_name),
            'орг_фио_кратко': self.format_full_name(*organization.ceo_name.split()),
            'основание': organization.ceo_footing,
            'орг_кпп': organization.kpp,
            'орг_рс': organization.r_s,
            'орг_банк': organization.bank.upper(),
            'орг_бик': organization.bik,
            'орг_счет': organization.k_s,
            'орг_адрес': organization.address,
            'орг_email': organization.email,
            'год': str(now.year)[2::],
            'раб_дни': self.get_workday_phrase(timeframe),
            'дни': timeframe,
            'число': day_of_month,
            'месяц': month_num,
            'месяц_назв': month_name,
            'номер_договора': c_number
        }

        # Полный путь к шаблону
        file_path = os.path.join(settings.BASE_DIR, 'media/contracts/contract.docx')
        doc = Document(file_path)

        # Заменяем метки в документе
        self._replace_placeholders(doc, data)

        # Определяем путь, по которому будет сохраняться новый документ
        num = c_number.replace("/", '')
        new_file_path = os.path.join(settings.MEDIA_ROOT, f'contracts/договор_{num}.docx')
        os.makedirs(os.path.dirname(new_file_path), exist_ok=True)

        # Сохраняем изменённый документ
        doc.save(new_file_path)

        return new_file_path, c_number

    def _replace_placeholders(self, doc, data):
        """Замена плейсхолдеров в документе Word"""

        def replace_in_paragraphs(paragraphs):
            for paragraph in paragraphs:
                for key, value in data.items():
                    if f'{{{key}}}' in paragraph.text or f'[[{key}]]' in paragraph.text:
                        paragraph.text = paragraph.text.replace(f'{{{key}}}', str(value))
                        paragraph.text = paragraph.text.replace(f'[[{key}]]', str(value))

        # Заменяем метки в главных параграфах
        replace_in_paragraphs(doc.paragraphs)

        # Заменяем метки в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_paragraphs(cell.paragraphs)