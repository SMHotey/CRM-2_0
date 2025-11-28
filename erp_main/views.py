import json
import os
import pymorphy3
from django.conf import settings
from django.contrib.auth import login, authenticate
from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator
from django.db import IntegrityError
from django.db.models import Sum, Q
from django.utils import timezone
from django.http import JsonResponse, HttpResponseBadRequest
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST, require_http_methods
from django.views.generic import FormView, CreateView, ListView, DetailView, UpdateView

from erp_main.views.permissions import get_user_role_from_request
from .models import Order, OrderItem, Organization, Invoice, InternalLegalEntity, GlassInfo, OrderChangeHistory, Contract, \
    Shipment
from .forms import OrderForm, InvoiceForm, OrderFileForm, ShipmentForm, InternalLegalEntityForm
import logging
from docx import Document
from datetime import datetime, timedelta, time
from collections import Counter
from django.shortcuts import redirect, get_object_or_404, render
from django.urls import reverse
from openpyxl import load_workbook
from django.contrib.auth.mixins import LoginRequiredMixin
import re
from django.utils.http import url_has_allowed_host_and_scheme
from django.contrib.auth.models import User
import json

logger = logging.getLogger(__name__)


@require_http_methods(["GET", "POST"])
def custom_login(request):
    # Обработка GET-запроса
    if request.method == 'GET':
        next_url = request.GET.get('next', '')  # Не устанавливаем дефолтное значение
        return render(request, 'registration/login.html', {
            'next': next_url,
        })

    # Обработка POST-запроса
    username = request.POST.get('username')
    password = request.POST.get('password')
    next_url = request.POST.get('next', '')  # Получаем без дефолтного значения

    user = authenticate(request, username=username, password=password)

    if user is None:
        # Неверные данные
        return render(request, 'registration/login.html', {
            'error': 'Неверные имя пользователя или пароль',
            'next': next_url,
        })

    if not user.is_active:
        # Пользователь неактивен
        return render(request, 'registration/login.html', {
            'error': 'Ваш аккаунт деактивирован',
            'next': next_url,
        })

    login(request, user)

    # Безопасный редирект
    if next_url and url_has_allowed_host_and_scheme(url=next_url, allowed_hosts={request.get_host()}):
        return redirect(next_url)
    return redirect(reverse('index'))  # Используем reverse вместо строки


@login_required  # Декоратор для проверки аутентификации пользователя
def index(request):
    start_date = request.GET.get('startDate', f'{datetime.now().year}-01-01')
    end_date = request.GET.get('endDate', timezone.now().strftime('%Y-%m-%d'))
    invoice_status = request.GET.get('invoice_status', '')
    order_status = request.GET.get('order_status', '')

    # Преобразуем строки в даты
    start_date_obj = datetime.strptime(start_date, '%Y-%m-%d').date()
    end_date_obj = datetime.strptime(end_date, '%Y-%m-%d').date() + timedelta(days=1)

    # Фильтруем организации по пользователю
    user_orgs = Organization.objects.filter(user=request.user)

    # Базовые запросы с фильтрацией по дате
    invoice_query = Invoice.objects.filter(
        organization__in=user_orgs,
        date__range=[start_date_obj, end_date_obj]
    )

    order_query = Order.objects.filter(
        invoice__organization__in=user_orgs,
        created_at__range=[start_date_obj, end_date_obj]
    )

    # Применяем фильтры статусов
    if invoice_status == 'paid':
        invoice_query = invoice_query.filter(is_paid=True)
    elif invoice_status == 'unpaid':
        invoice_query = invoice_query.filter(is_paid=False)

    if order_status:
        # Для заказов фильтрация идет по статусам позиций
        order_query = order_query.filter(
            items__p_status=order_status
        ).distinct()

    # Получаем данные
    user_invoices = invoice_query.order_by('-date')
    user_orders = order_query.order_by('-created_at')
    user_role = get_user_role_from_request(request)

    # Подготовка контекста
    context = {
        'current_year': datetime.now().year,
        'start_date': start_date,
        'end_date': end_date,
        'invoice_status': invoice_status,
        'order_status': order_status,
        'orgs_count': user_orgs.count(),
        'orders_count': user_orders.count(),
        'invoices_count': user_invoices.count(),
        'total_invoices_amount': user_invoices.aggregate(total=Sum('amount'))['total'] or 0,
        'invoices': user_invoices,
        'orders': user_orders,
        'organizations': user_orgs,
        'user_role': user_role,
    }

    return render(request, 'index.html', context)


# class OrganizationCreateView(LoginRequiredMixin, CreateView):
#     model = Organization
#     form_class = OrganizationForm
#     template_name = 'organization_add.html'
#
#     def form_valid(self, form):
#         organization = form.save(commit=False)
#         organization.user = self.request.user
#         organization.save()
#         return redirect('organization_list')
#
#
# class OrganizationUpdateView(LoginRequiredMixin, UpdateView):
#     model = Organization
#     form_class = OrganizationForm
#     template_name = 'organization_edit.html'
#     context_object_name = 'organization'
#
#     def get_success_url(self):
#         return reverse('organization_detail', kwargs={'pk': self.object.pk})
#
#     def form_valid(self, form):
#         organization = form.save(commit=False)
#         type_ = form.cleaned_data.get('type')
#
#         if type_ == 'organization':
#             organization.name_fl = None
#             organization.phone_number = None
#         elif type_ == 'individual':
#             organization.name = None
#             organization.inn = None
#             organization.ogrn = None
#             organization.kpp = None
#             organization.r_s = None
#             organization.bank = None
#             organization.bik = None
#             organization.k_s = None
#             organization.ceo_title = None
#
#         organization.save()
#         return super().form_valid(form)
#
#
# class OrganizationListView(LoginRequiredMixin, ListView):
#     model = Organization
#     template_name = 'organization_list.html'
#     context_object_name = 'organizations'
#
#     def get_queryset(self):
#         user = self.request.user
#         if user.is_superuser:
#             return Organization.objects.all()
#         return Organization.objects.filter(user=user)
#
#
# class OrganizationDetailView(LoginRequiredMixin, DetailView):
#     model = Organization
#     template_name = 'organization_detail.html'
#     context_object_name = 'organization'
#
#     def get_context_data(self, **kwargs):
#         context = super().get_context_data(**kwargs)
#         context['legal_entities'] = LegalEntity.objects.all()
#         return context


class OrderUploadView(LoginRequiredMixin, FormView):
    template_name = 'order_upload.html'
    form_class = OrderForm

    def get_context_data(self, **kwargs):

        context = super().get_context_data(**kwargs)
        if self.request.user.is_superuser:
            context['organizations'] = Organization.objects.all()
        else:
            context['organizations'] = Organization.objects.filter(user=self.request.user)

        context['legal_entities'] = InternalLegalEntity.objects.all()

        return context

    def get_form_kwargs(self):
        """Передаем текущего пользователя в форму"""
        kwargs = super().get_form_kwargs()
        kwargs['user'] = self.request.user
        return kwargs

    def form_valid(self, form):
        uploaded_file = form.cleaned_data.get('order_file')
        order_id = self.kwargs.get('order_id')

        # Для нового заказа файл обязателен
        if not order_id and not uploaded_file:
            form.add_error('order_file', 'Пожалуйста, загрузите файл.')
            return self.form_invalid(form)

        if order_id:
            # Извлечение существующего заказа из базы
            order = get_object_or_404(Order, pk=order_id)
            old_file = order.order_file  # старый файл заказа для сохранения в истории изменений
            is_update = True
        else:
            # Создание нового заказа на основе данных из формы
            order = Order()
            is_update = False
            order.invoice = form.cleaned_data.get('invoice')  # Извлекаем счёт из формы
            order.due_date = form.cleaned_data.get('due_date')  # Извлекаем дату готовности
            order.comment = form.cleaned_data.get('comment', '')  # Извлекаем комментарий

        # Обновление поля order_file только если файл был загружен
        if uploaded_file:
            order.order_file = uploaded_file  # Загружаем файл, если он был загружен

            # Проверка файла
            try:
                wb = load_workbook(uploaded_file)
                if not self._check_header(wb.active):
                    form.add_error('order_file', 'Выберите правильный файл заказа')
                    return self.form_invalid(form)
            except Exception as e:
                form.add_error('order_file', 'Ошибка загрузки файла: ' + str(e))
                return self.form_invalid(form)

        # Сохраняем объект Order
        order.save()

        # Обработка загрузки файла и обновление позиций заказа
        if uploaded_file:  # Обработка файла только если он был загружен
            wb = load_workbook(uploaded_file)
            new_positions = self._process_file(wb.active)

            if is_update:
                self._update_order_items(order, new_positions, old_file)  # Обновление позиций для существующего заказа
            else:
                self._create_order_items(order, new_positions)  # Создание новых позиций для нового заказа

        return redirect(reverse('order_detail', args=[order.id]))

    def _render_form_with_context(self, form):
        organizations = Organization.objects.all()
        return self.render_to_response(self.get_context_data(form=form, organizations=organizations))

    def _check_header(self, sheet):
        return sheet.cell(row=1, column=3).value == "Бланк №"

    def _process_file(self, sheet):
        cur_row, cur_column = 9, 15
        while sheet.cell(row=cur_row, column=cur_column).value != 'шт.':
            cur_row += 1
        max_row = cur_row

        seq = [1, 2, 3, 4, 5, 6, 9, 10, 11, 12, 13, 14, 15, 7, 8]
        positions = []
        line = []

        for row in range(8, max_row):
            if sheet.cell(row=row, column=2).value:
                if line:
                    positions.append(line)
                line = [sheet.cell(row=row, column=column).value for column in seq]
            else:
                line.extend([sheet.cell(row=row, column=7).value, sheet.cell(row=row, column=8).value])

        if line:
            positions.append(line)

        return positions

    def _update_order_items(self, order, new_positions, old_file):
        current_items = {item.position_num: item for item in
                         OrderItem.objects.filter(order=order).prefetch_related('glasses')}
        current_date = datetime.now().strftime('%d.%m.%Y')
        changes_made = []

        for data in new_positions:
            n_num = str(data[0])
            name = data[1]
            n_kind = self._get_kind(name)
            n_type = self._get_type(name)
            n_construction = 'NK' if re.search('-м', name.lower()) else 'SK'

            new_item_data = {
                'p_kind': n_kind,
                'p_type': n_type,
                'p_construction': n_construction,
                'p_height': data[2],
                'p_width': data[3],
                'p_active_trim': data[4],
                'p_open': data[5],
                'p_platband': data[6],
                'p_furniture': data[7],
                'p_door_closer': data[8],
                'p_step': data[9],
                'p_ral': str(data[10]),
                'p_quantity': data[11],
                'p_comment': data[12],
                'p_glass': self._count_glass(data[13:]),
            }

            if n_num in current_items:
                first = 0
                current_item = current_items[n_num]
                for field, new_value in new_item_data.items():
                    old_value = getattr(current_item, field, None)

                    # Handle glass changes separately
                    if field == 'p_glass':
                        # Get existing glass info from database
                        old_glass = {(g.height, g.width): g.quantity for g in current_item.glasses.all()}
                        new_glass = new_value

                        # Check if glass has changed
                        if old_glass != new_glass:
                            if first == 0:
                                changes_made.append(f'<br> поз. {n_num}: изменено стекло ')
                                first = 1

                            # Format old glass info
                            old_glass_str = ", ".join(
                                [f"{h}x{w} ({q} шт.)" for (h, w), q in old_glass.items()]) if old_glass else "нет"
                            # Format new glass info
                            new_glass_str = ", ".join(
                                [f"{h}x{w} ({q} шт.)" for (h, w), q in new_glass.items()]) if new_glass else "нет"

                            changes_made.append(f"с \"{old_glass_str}\" на \"{new_glass_str}\";")

                            # Update p_glass field
                            setattr(current_item, 'p_glass', str(new_glass))
                    else:
                        if str(old_value) != str(new_value):
                            field_name = OrderItem._meta.get_field(field).verbose_name
                            if old_value:
                                if first == 0:
                                    changes_made.append(
                                        f'<br> поз. {n_num}:  {field_name} с "{old_value}" на "{new_value}"; ')
                                    first = 1
                                else:
                                    changes_made.append(f'{field_name} с "{old_value}" на "{new_value}";')
                            else:
                                changes_made.append(f'поз. {n_num}: добавлен {field_name} "{new_value}";')
                            setattr(current_item, field, new_value)

                if changes_made:
                    # Delete existing glass info before saving new one
                    current_item.glasses.all().delete()
                    current_item.save()

                    # Save new glass info
                    for (height, width), quantity in new_item_data['p_glass'].items():
                        if height and width:  # Only create records for valid dimensions
                            GlassInfo.objects.create(
                                order_items=current_item,
                                height=height,
                                width=width,
                                quantity=quantity
                            )
            else:
                new_item = OrderItem(order=order, position_num=n_num, **new_item_data)
                new_item.p_status = 'in_query'
                new_item.save()

                # Save glass info for new item
                for (height, width), quantity in new_item_data['p_glass'].items():
                    if height and width:  # Only create records for valid dimensions
                        GlassInfo.objects.create(
                            order_items=new_item,
                            height=height,
                            width=width,
                            quantity=quantity
                        )

        if changes_made:
            comment = str(changes_made).replace('[', '').replace(']', '').replace("'", "").replace(",", "").strip()[5::]
            order.save()
            add_changes = OrderChangeHistory(order=order, order_file=old_file, changed_by=self.request.user,
                                             comment=comment)
            add_changes.save()

            if order.order_file:
                try:
                    wb = load_workbook(order.order_file.path)
                    sheet = wb.active
                    sheet['K3'] = comment.replace('<br>', '')
                    wb.save(order.order_file.path)
                except Exception as e:
                    print(f"Ошибка при записи комментария в файл: {e}")

    def _create_order_items(self, order, new_positions):
        for data in new_positions:
            n_num = data[0]
            name = data[1]
            n_kind = self._get_kind(name)
            n_type = self._get_type(name)
            n_construction = 'NK' if re.search('-м', name.lower()) else 'SK'

            new_item_data = {
                'p_kind': n_kind,
                'p_type': n_type,
                'p_construction': n_construction,
                'p_height': data[2],
                'p_width': data[3],
                'p_active_trim': data[4],
                'p_open': data[5],
                'p_platband': data[6],
                'p_furniture': data[7],
                'p_door_closer': data[8],
                'p_step': data[9],
                'p_ral': data[10],
                'p_quantity': data[11],
                'p_comment': data[12],
                'p_glass': self._count_glass(data[13:]),
            }

            new_item = OrderItem(order=order, position_num=n_num, **new_item_data)
            new_item.p_status = 'in_query'
            new_item.save()  # Сохраняем новый объект

            # Сохраняем информацию о стеклах
            self._save_glass_info(new_item, new_item_data['p_glass'])

    def _get_kind(self, name):
        kind_mapping = {
            'дверь': 'door', 'люк': 'hatch', 'ворота': 'gate',
            'калитка': 'door', 'фрамуга': 'transom'
        }
        return next((value for key, value in kind_mapping.items() if re.search(key, name, re.IGNORECASE)), None)

    def _get_type(self, name):
        type_mapping = {
            'ei-60': 'ei-60', 'eis-60': 'eis-60', 'eiws-60': 'eiws-60',
            'тех': 'tech', 'ревиз': 'revision'
        }
        return next((value for key, value in type_mapping.items() if re.search(key, name, re.IGNORECASE)), None)

    def _save_glass_info(self, new_item, counted_glass):
        """Сохраняет информацию о стеклах для нового элемента заказа."""
        for (height, width), quantity in counted_glass.items():
            if height and width:  # Проверяем, что высота и ширина заданы
                # Создаем новый объект GlassInfo, который содержит внешний ключ к OrderItem
                new_glass = GlassInfo(height=height, width=width, quantity=quantity, order_items=new_item)
                new_glass.save()  # Сохраняем объект GlassInfo

    def _count_glass(self, glass_data):
        counted_glass = dict(Counter(list(zip(glass_data[::2], glass_data[1::2]))))
        counted_glass.pop((None, None), None)
        return counted_glass

    def form_invalid(self, form):
        organizations = Organization.objects.all()
        return self.render_to_response(self.get_context_data(form=form, organizations=organizations))


def glass(request):
    return render(request, 'glass_info.html')


@login_required(login_url='login')
def invoice_add(request):
    if request.method == 'POST':
        referer_url = request.META.get('HTTP_REFERER')
        expected_url = request.build_absolute_uri(reverse('invoice_add'))
        form = InvoiceForm(request.user, request.POST, request.FILES)
        if form.is_valid():
            try:
                invoice = form.save()
                if referer_url != expected_url:
                    return JsonResponse({
                        'success': True,
                        'invoice_id': invoice.id,
                        'invoice_number': form.cleaned_data['number'],
                        'message': 'Счет добавлен'})
                else:
                    return redirect(invoices_list)

            except IntegrityError:
                return JsonResponse({
                    'success': False,
                    'error': 'Запись с такими значениями полей уже существует.'
                }, status=400)

        error_messages = form.errors.as_json()
        return JsonResponse({'success': False, 'error': error_messages}, status=400)

    form = InvoiceForm(request.user)  # Создание экземпляра формы
    return render(request, 'invoice_add.html', {'form': form})  # Возвращаем шаблон с формой


def orders_list(request):
    orders = []
    source = request.GET.get('source')
    if request.user.is_staff:
        orders = Order.objects.all().order_by('-id')
    elif Order.objects.filter(invoice__organization__user=request.user):
        orders = Order.objects.all().order_by('-id')
#        orders = Order.objects.filter(invoice__organization__user=request.user).order_by('-id')

    if source:
        orders = Order.objects.filter(invoice__organization=source).order_by('-id')

    paginator = Paginator(orders, 20)  # Создаем пагинатор
    page_number = request.GET.get('page')  # Получаем номер страницы из GET параметров
    orders = paginator.get_page(page_number)  # Получаем заказы для текущей страницы

    return render(request, 'orders_list.html', {'orders': orders})


@login_required(login_url='login')
def order_detail(request, order_id):
    order = get_object_or_404(Order, id=order_id)
    changes = order.changes.all()

    # Отфильтрованные OrderItem, где статус не равен 'changed'
    filtered_items = order.items.exclude(p_status__in=['changed', ])  #фильтрация по активным позициям

    if request.method == 'POST':
        form = OrderFileForm(request.POST, request.FILES)
        if form.is_valid():
            order.order_file = form.cleaned_data['order_file']  # Замените файл заказа
            order.save()
            return redirect('order_detail', order_id=order.id)  # Ссылаемся на order_id

    context = {
        'order': order,
        'filtered_items': filtered_items,
        'changes': changes,
    }

    return render(request, 'order_detail.html', context)


@login_required(login_url='login')
def invoice_detail(request, pk):
    invoice = get_object_or_404(Invoice, id=pk)
    orders = invoice.invoice.all()

    if request.method == 'POST':
        form = InvoiceForm(user=request.user, data=request.POST, files=request.FILES, instance=invoice)
        if form.is_valid():
            invoice = form.save()
            file_url = invoice.invoice_file.url if invoice.invoice_file else None

            return JsonResponse({
                'success': True,
                'file_url': file_url,
                'redirect_url': reverse('invoice_detail', args=[invoice.id])
            })
        else:
            error_messages = form.errors.as_json()
            return JsonResponse({'success': False, 'error': error_messages}, status=400)

    else:
        form = InvoiceForm(user=request.user, instance=invoice)

    return render(request, 'invoice_detail.html', {'invoice': invoice, 'orders': orders, 'form': form})


@login_required(login_url='login')
def invoices_list(request):
    search_query = request.GET.get('search', '')
    selected_legal_entity_id = request.GET.get('legal_entity', None)
    sort_by = request.GET.get('sort', 'id')
    direction = request.GET.get('direction', 'desc')
    source = request.GET.get('source')

    # Начальное значение для queryset
    if request.user.is_staff:
        invoices = Invoice.objects.all()
    else:
        invoices = Invoice.objects.filter(organization__user=request.user)

    # Фильтрация по поисковому запросу
    if search_query:
        invoices = invoices.filter(number__icontains=search_query) | invoices.filter(
            organization__name__icontains=search_query)

    # Фильтрация по выбранному юридическому лицу
    if selected_legal_entity_id:
        # Выберите только если selected_legal_entity_id является числом
        try:
            selected_legal_entity_id = int(selected_legal_entity_id)  # Приводим к числу
            invoices = invoices.filter(legal_entity_id=selected_legal_entity_id)
        except ValueError:
            pass  # Игнорируем, если не удалось привести к числу

    # Сортировка
    if sort_by == 'number':
        invoices = invoices.order_by(f"{'-' if direction == 'desc' else ''}id")
    else:
        invoices = invoices.order_by('-id')  # Сортировка по умолчанию

    if source:
        organization = Organization.objects.filter(id=source).first()
        invoices = invoices.filter(organization=organization)

    # Пагинация
    paginator = Paginator(invoices, 10)
    page_number = request.GET.get('page')
    invoices_page = paginator.get_page(page_number)

    # Получение всех юридических лиц для отображения в выпадающем списке
    legal_entities = InternalLegalEntity.objects.all()

    return render(request, 'invoices_list.html', {
        'invoices': invoices_page,
        'legal_entities': legal_entities,
        'selected_legal_entity_id': selected_legal_entity_id,
        'request': request,
    })


def update_order_item_status(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            updates = data.get('updates', {})

            for item_id in updates.keys():
                order_item = get_object_or_404(OrderItem, id=item_id)
                new_data = updates[item_id]
                order_item.p_status = new_data['status']
                order_item.workshop = new_data['workshop']
                if order_item.workshop == '2' and new_data['path'] != 'order_detail':
                    order_item.p_status = 'stopped'
                if ((order_item.p_status == 'stopped' or order_item.p_status == 'canceled') and
                        new_data['path'] != 'order_detail'):
                    order_item.workshop = '2'
                order_item.save()

            return JsonResponse({'status': 'success', 'message': 'Информация обновлена'})

        except json.JSONDecodeError:
            return JsonResponse({'status': 'error', 'message': 'Неверный формат'}, status=400)
        except Exception as e:
            logger.exception("Error updating order items' statuses")
            return JsonResponse({'status': 'error', 'message': 'An error occurred while processing your request.'},
                                status=500)

    return JsonResponse({'status': 'error', 'message': 'Invalid request method.'}, status=405)


def create_legal_entity(request):
    if request.method == 'POST':
        form = InternalLegalEntityForm(request.user, request.POST)
        if form.is_valid():
            form.save()  # Сохраните объект LegalEntity
            return redirect('organization_list')  # Перенаправление на страницу успеха
    else:
        form = InternalLegalEntityForm(request.user)

    return render(request, 'legal_entity_form.html', {'form': form})


def create_contract(request, pk):
    timeframe = 21
    now = datetime.now()
    months_ru = [
        "января", "февраля", "марта", "апреля", "мая", "июня",
        "июля", "августа", "сентября", "октября", "ноября", "декабря"
    ]
    day_of_month = now.day
    month_num = now.month
    month_name = months_ru[month_num - 1]

    def genitive_case(word):
        if not word:  # Если слово пустое, возвращаем пустую строку
            return ""
        morph = pymorphy3.MorphAnalyzer()
        try:
            words = word.split()  # Разделяем словосочетание на отдельные слова
            modified_words = []  # Список для измененных слов

            for w in words:
                parsed_word = morph.parse(w)[0]  # Анализируем каждое слово
                if "NOUN" in parsed_word.tag or "ADJF" in parsed_word.tag:  # Если это существительное или прилагательное
                    genitive = parsed_word.inflect({"gent"}).word  # Склоняем в родительный падеж
                    modified_words.append(genitive)  # Добавляем в список
                else:
                    modified_words.append(w)  # Если это не существительное или прилагательное, добавляем исходное слово

            # Проверяем, состоит ли словосочетание из трех слов (например, ФИО)
            if len(words) == 3:
                if 'уляшва' in modified_words:
                    modified_words[0] = 'уляшова'
                capitalized_words = [word.capitalize() for word in modified_words]  # Капитализируем каждое слово
                return " ".join(capitalized_words)  # Соединяем слова обратно в строку
            else:
                return " ".join(modified_words)  # Возвращаем без изменений в регистре

        except Exception as e:
            print(f"Ошибка при склонении слова '{word}': {e}")
            return word  # В случае ошибки возвращаем исходное слово

    def format_full_name(surname, name, patronymic):
        # Преобразование в формат "Фамилия И. О."
        formatted_name = f"{surname.capitalize()} {name[0].upper()}. {patronymic[0].upper()}."
        return formatted_name

    def get_workday_phrase(number):
        # Инициализация морфологического анализатора
        morph = pymorphy3.MorphAnalyzer()

        # Определение формы слова "рабочий"
        if number % 10 == 1 and number % 100 != 11:
            workday_form = morph.parse('рабочий')[0].inflect(
                {'nomn', 'sing'}).word  # Именительный падеж, единственное число
        elif number % 10 in [2, 3, 4] and not (number % 100 in [12, 13, 14]):
            workday_form = morph.parse('рабочий')[0].inflect(
                {'gent', 'plur'}).word  # Родительный падеж, множественное число
        else:
            workday_form = morph.parse('рабочий')[0].inflect(
                {'gent', 'plur'}).word  # Родительный падеж, множественное число

        # Определение формы слова "день"
        if number % 10 == 1 and number % 100 != 11:
            day_form = 'день'  # 1, 21
        elif number % 10 in [2, 3, 4] and not (number % 100 in [12, 13, 14]):
            day_form = 'дня'  # 2, 3, 4, 22, 23, 24
        else:
            day_form = 'дней'  # 5-20, 25-31

        # Формируем строку
        return f"{number} {workday_form} {day_form}"

    if request.method == 'POST':
        legal_entity_id = request.POST.get('legal_entity')
        legal_entity = get_object_or_404(InternalLegalEntity, pk=legal_entity_id)
        organization = get_object_or_404(Organization, pk=pk)
        c_number = f'{day_of_month}/{month_num}/{str(datetime.now().year)[2::]}/36/{organization.inn[-4:]}'

        data = {
            'юл': legal_entity.name.upper(),
            'юл_огрн': legal_entity.ogrn,
            'юл_инн': legal_entity.inn,
            'юл_должность': legal_entity.ceo_title,
            'юл_должность_рп': genitive_case(legal_entity.ceo_title),
            'юл_фио': legal_entity.ceo_name,
            'юл_фио_рп': genitive_case(legal_entity.ceo_name),
            'юл_фио_кратко': format_full_name(*legal_entity.ceo_name.split()),
            'юл_кпп': legal_entity.kpp,
            'юл_рс': legal_entity.r_s,
            'юл_банк': legal_entity.bank,
            'юл_бик': legal_entity.bik,
            'юл_корс': legal_entity.k_s,
            'юл_адрес': legal_entity.address,
            'юл_email': legal_entity.email,
            'орг': organization.name.upper(),
            'орг_огрн': organization.ogrn,
            'орг_инн': organization.inn,
            'инн_4': organization.inn[-4:],
            'орг_должность': organization.ceo_title,
            'орг_должность_рп': genitive_case(organization.ceo_title),
            'орг_фио': organization.ceo_name,
            'орг_фио_рп': genitive_case(organization.ceo_name),
            'орг_фио_кратко': format_full_name(*organization.ceo_name.split()),
            'основание': organization.ceo_footing,
            'орг_кпп': organization.kpp,
            'орг_рс': organization.r_s,
            'орг_банк': organization.bank.upper(),
            'орг_бик': organization.bik,
            'орг_счет': organization.k_s,
            'орг_адрес': organization.address,
            'орг_email': organization.email,
            'год': str(datetime.now().year)[2::],
            'раб_дни': get_workday_phrase(timeframe),
            'дни': timeframe,
            'число': day_of_month,
            'месяц': month_num,
            'месяц_назв': month_name,
            'номер_договора': c_number
        }

        # Полный путь к шаблону
        file_path = os.path.join(settings.BASE_DIR, 'media/contracts/contract.docx')
        doc = Document(file_path)

        # Функция для замены меток в параграфах
        def replace_in_paragraphs(paragraphs, data):
            for paragraph in paragraphs:
                for key, value in data.items():
                    if f'{{{key}}}' in paragraph.text or f'[[{key}]]' in paragraph.text:
                        paragraph.text = paragraph.text.replace(f'{{{key}}}', str(value))
                        paragraph.text = paragraph.text.replace(f'[[{key}]]', str(value))

        # Заменяем метки в главных параграфах
        replace_in_paragraphs(doc.paragraphs, data)

        # Заменяем метки в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_paragraphs(cell.paragraphs, data)

        # Определяем путь, по которому будет сохраняться новый документ
        num = c_number.replace("/", '')
        new_file_path = os.path.join(settings.MEDIA_ROOT, f'contracts/договор_{num}.docx')
        os.makedirs(os.path.dirname(new_file_path), exist_ok=True)  # Создаем директорию, если её нет

        # Сохраняем изменённый документ
        doc.save(new_file_path)

        # URL к новому файлу

        new_contract_url = os.path.join(settings.MEDIA_URL, f'contracts/договор_{num}.docx')
        legal_entities = InternalLegalEntity.objects.all()

        new_contract = Contract(number=num, organization=organization, legal_entity=legal_entities, days=timeframe,
                                file=new_file_path)
        new_contract.save()

        return render(request, 'organizations/organization_detail.html', {
            'organization': organization,
            'new_contract_url': new_contract_url,
            'legal_entities': legal_entities,
        })



def glass_info(request, pk=''):
    if pk == '':
        orders = Order.objects.all()
    else:
        orders = [get_object_or_404(Order, pk=pk)]

    return render(request, 'glass_info.html', {'orders': orders,
                                               'glass_options': GlassInfo.OPTIONS_CHOICE,
                                               'glass_status': GlassInfo.GLASS_STATUS_CHOICE})


@require_POST
def update_glass_status(request, glass_id):
    glass = get_object_or_404(GlassInfo, id=glass_id)
    new_status = request.POST.get('status')
    if new_status in dict(GlassInfo.GLASS_STATUS_CHOICE).keys():
        glass.status = new_status
        glass.save()
    return redirect('glass_info')  # Перенаправление обратно на страницу с заказами


def update_workshop(request, order_id):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            workshop_value = data.get('workshop')

            # Проверяем существование заказа
            if not OrderItem.objects.filter(order_id=order_id).exists():
                return JsonResponse({'success': False, 'error': 'Order not found'}, status=404)

            new_status = ''
            status_list = {'product': '"запущен"', 'stopped': '"остановлен"', 'ready': '"готов"'}

            if workshop_value in ['1', '3']:
                OrderItem.objects.filter(order_id=order_id).update(workshop=workshop_value)
                new_status = 'product'
            elif workshop_value == '2':
                OrderItem.objects.filter(order_id=order_id).update(workshop=workshop_value)
                new_status = 'stopped'
            elif workshop_value == '4':
                new_status = 'ready'
            else:
                return JsonResponse({'success': False, 'error': 'Invalid workshop value'}, status=400)

            OrderItem.objects.filter(order_id=order_id).update(p_status=new_status)

            add_changes = OrderChangeHistory(
                order_id=order_id,
                changed_by=request.user,
                comment='статус заказа изменен на ' + status_list[new_status]
            )
            add_changes.save()

            return JsonResponse({'success': True})
        except json.JSONDecodeError:
            return JsonResponse({'success': False, 'error': 'Invalid JSON'}, status=400)
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)}, status=500)

    return JsonResponse({'success': False, 'error': 'Invalid request method'}, status=400)


def make_passport(self):
    pass


def calculate(self: OrderItem):
    pass


# views.py
@csrf_exempt
@require_http_methods(["POST"])
def save_shipment(request):
    try:
        data = request.POST
        shipment_id = data.get('shipment_id')

        if not request.user.is_authenticated:
            return JsonResponse({'status': 'error', 'message': 'Требуется авторизация'}, status=403)

        if shipment_id:
            shipment = get_object_or_404(Shipment, id=shipment_id)
            if shipment.user != request.user and not request.user.is_superuser:
                return JsonResponse({'status': 'error', 'message': 'Нет прав на редактирование'}, status=403)
        else:
            shipment = Shipment(
                user=request.user,
                date=data.get('date'),
                time=data.get('time'),
                workshop=data.get('workshop')
            )

        # Обновляем основную информацию
        if data.get('order'):
            try:
                shipment.order = Order.objects.get(pk=data.get('order'))
            except Order.DoesNotExist:
                return JsonResponse({'status': 'error', 'message': 'Указанный заказ не существует'}, status=400)

        # Обновляем JSON-поля
        order_items = shipment.order_items or {}
        order_items['type'] = data.get('order_type', '')
        shipment.order_items = order_items

        car_info = shipment.car_info or {}
        car_info.update({
            'brand': data.get('car_brand', ''),
            'number': data.get('car_number', '')
        })
        shipment.car_info = car_info

        driver_info = shipment.driver_info or {}
        driver_info.update({
            'comments': data.get('comments', ''),
            'shipment_mark': data.get('shipment_mark', '')
        })
        shipment.driver_info = driver_info

        shipment.address = data.get('address', '')
        shipment.save()

        return JsonResponse({
            'status': 'success',
            'shipment_id': shipment.id,
            'order': shipment.order.pk if shipment.order else '',
            'order_type': shipment.order_items.get('type', ''),
            'car_brand': shipment.car_info.get('brand', ''),
            'car_number': shipment.car_info.get('number', ''),
            'address': shipment.address,
            'comments': shipment.driver_info.get('comments', ''),
            'shipment_mark': shipment.driver_info.get('shipment_mark', '')
        })

    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@require_http_methods(["GET"])
def shipment_detail(request, workshop, date):
    # Преобразуем строку даты в объект date
    try:
        date_obj = datetime.strptime(date, '%Y-%m-%d').date()
    except ValueError:
        return HttpResponseBadRequest("Неверный формат даты")

    # Получаем отгрузки для указанного цеха и даты
    shipments = Shipment.objects.filter(workshop=workshop, date=date_obj).order_by('time')

    # Создаем словарь отгрузок по времени для удобного доступа в шаблоне
    shipments_dict = {shipment.time: shipment for shipment in shipments}

    # Генерируем список времен для отображения в таблице
    times = [time(hour, minute) for hour in range(9, 18) for minute in [0, 30]]

    # Получаем список всех неотгруженных заказов
    filtered_items = OrderItem.objects.filter(p_status__in=['product', 'ready'])
    orders = Order.objects.filter(items__in=filtered_items).distinct()
    orders_list = [{'pk': order.pk, 'created_at': order.created_at.strftime('%Y-%m-%d %H:%M:%S')} for order in orders]

    context = {
        'workshop': workshop,
        'date': date_obj,
        'shipments': shipments_dict,  # Используем словарь вместо QuerySet
        'times': times,
        'orders': orders_list,
    }
    return render(request, 'shipment_detail.html', context)


@csrf_exempt
@require_http_methods(["POST"])
def delete_shipment(request, shipment_id):
    try:
        if not request.user.is_authenticated:
            return JsonResponse({'status': 'error', 'message': 'Требуется авторизация'}, status=403)

        shipment = get_object_or_404(Shipment, id=shipment_id)

        # Проверка прав (только создатель или админ)
        if shipment.user != request.user and not request.user.is_superuser:
            return JsonResponse({'status': 'error', 'message': 'Нет прав на удаление'}, status=403)

        shipment.delete()

        return JsonResponse({
            'status': 'success',
            'message': 'Отгрузка успешно удалена'
        })
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


def calendar_view(request):
    # Получаем текущую дату
    today = timezone.now().date()

    # Получаем год и месяц из запроса (если они переданы)
    year = int(request.GET.get('year', today.year))
    month = int(request.GET.get('month', today.month))

    # Создаем объект даты для текущего месяца
    current_date = datetime(year, month, 1).date()

    # Вычисляем даты для предыдущего и следующего месяца
    prev_month = (current_date - timedelta(days=1)).replace(day=1)
    next_month = (current_date + timedelta(days=31)).replace(day=1)

    # Контекст для передачи в шаблон
    context = {
        'current_date': current_date,
        'prev_month': prev_month,
        'next_month': next_month,
        'workshop1': 1,  # ID цеха №1
        'workshop3': 3,  # ID цеха №3
    }

    return render(request, 'calendar.html', context)


# Добавьте в erp_main/views.py

def debug_users(request):
    """Временный endpoint для диагностики"""
    users = User.objects.all()
    user_data = []

    for user in users:
        user_data.append({
            'id': user.id,
            'username': user.username,
            'first_name': user.first_name,
            'last_name': user.last_name,
            'email': user.email,
        })

    return JsonResponse({
        'total_users': users.count(),
        'users': user_data,
        'debug': 'This is from debug endpoint'
    })



