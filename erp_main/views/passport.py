import logging
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_GET
import os
from datetime import datetime
from erp_main.models import Order, Certificate

from docxtpl import DocxTemplate
import io
from docx import Document
from docxcompose.composer import Composer



logger = logging.getLogger(__name__)


@require_GET
def check_nameplates(request):
    """Проверка наличия шильдов для всех позиций"""
    try:
        order_id = request.GET.get('order_id')
        if not order_id:
            return JsonResponse({'error': 'Не указан ID заказа'}, status=400)

        order = Order.objects.get(id=order_id)
        # Позиции, требующие шильдов
        target_types = ['ei-60', 'eis-60', 'eiws-60']
        items_requiring_nameplates = order.items.filter(p_type__in=target_types)

        missing_items = []
        has_all_nameplates = True

        for item in items_requiring_nameplates:
            if not item.nameplates.exists():
                has_all_nameplates = False
                missing_items.append({
                    'position_num': item.position_num,
                    'p_kind': item.get_p_kind_display(),
                    'p_type': item.get_p_type_display()
                })

        return JsonResponse({
            'has_all_nameplates': has_all_nameplates,
            'missing_items': missing_items
        })

    except Order.DoesNotExist:
        return JsonResponse({'error': 'Заказ не найден'}, status=404)
    except Exception as e:
        logger.error(f"Error in check_nameplates: {str(e)}")
        return JsonResponse({'error': f'Внутренняя ошибка сервера: {str(e)}'}, status=500)


@require_GET
def generate_passports(request):
    """Генерация паспортов для всех позиций заказа"""
    try:
        order_id = request.GET.get('order_id')
        if not order_id:
            return JsonResponse({'error': 'Не указан ID заказа'}, status=400)

        order = Order.objects.get(id=order_id)
        target_types = ['ei-60', 'eis-60', 'eiws-60']

        items_with_nameplates = order.items.filter(
            p_type__in=target_types,
            nameplates__isnull=False
        ).prefetch_related('nameplates', 'nameplates__certificate')

        if not items_with_nameplates.exists():
            return JsonResponse({'error': 'Нет позиций с шильдами для формирования паспортов'}, status=400)

        # Список для хранения всех сгенерированных документов
        generated_docs = []

        for item in items_with_nameplates:
            for nameplate in item.nameplates.all():
                certificate = nameplate.certificate

                if not certificate.passport_templates:
                    logger.warning(f"Нет шаблона паспорта для сертификата {certificate.id}")
                    continue

                if not os.path.exists(certificate.passport_templates.path):
                    logger.warning(f"Файл шаблона не существует: {certificate.passport_templates.path}")
                    continue

                # Определяем диапазон номеров
                if nameplate.end_value:
                    numbers = range(nameplate.first_value, nameplate.end_value + 1)
                else:
                    numbers = [nameplate.first_value]

                for number in numbers:
                    try:
                        # Загружаем шаблон
                        doc = DocxTemplate(certificate.passport_templates.path)

                        # Контекст для замены
                        context = {
                            'height': str(item.p_height or ''),
                            'width': str(item.p_width or ''),
                            'nameplate_number': str(number),
                            'issue_date': nameplate.issue_date.strftime('%d.%m.%Y') if nameplate.issue_date else '',
                            'certificate_number': nameplate.certificate.numbers if nameplate.certificate else '',
                            'position_number': item.position_num,
                            'product_type': item.get_p_type_display(),
                            'product_kind': item.get_p_kind_display(),
                            'quantity': item.p_quantity,
                        }

                        doc.render(context)

                        # Сохраняем документ в буфер
                        temp_buffer = io.BytesIO()
                        doc.save(temp_buffer)
                        temp_buffer.seek(0)

                        generated_docs.append({
                            'buffer': temp_buffer,
                            'item': item,
                            'nameplate_number': number
                        })

                    except Exception as e:
                        logger.error(f"Ошибка при обработке позиции {item.position_num}: {str(e)}", exc_info=True)
                        continue

        if not generated_docs:
            return JsonResponse({'error': 'Не удалось сгенерировать ни одного паспорта'}, status=400)

        # Объединяем все документы в один
        if len(generated_docs) == 1:
            # Если только один документ, возвращаем его как есть
            final_buffer = generated_docs[0]['buffer']
        else:
            # Объединяем несколько документов
            final_buffer = merge_documents(generated_docs)

        response = HttpResponse(
            final_buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="passports_order_{order_id}.docx"'
        return response

    except Exception as e:
        logger.error(f"Error in generate_passports: {str(e)}", exc_info=True)
        return JsonResponse({'error': f'Ошибка: {str(e)}'}, status=500)


def merge_documents(documents):
    """Объединение документов с принудительными разрывами страниц"""
    from docx.enum.text import WD_BREAK

    if not documents:
        return io.BytesIO()

    # Обрабатываем каждый документ - добавляем разрыв страницы в конец (кроме последнего)
    processed_docs = []

    for i, doc_data in enumerate(documents):
        # Загружаем документ из буфера
        doc_buffer = doc_data['buffer']
        doc_buffer.seek(0)
        temp_doc = Document(doc_buffer)

        # Добавляем разрыв страницы в конец документа (кроме последнего)
        if i < len(documents) - 1:
            last_paragraph = temp_doc.add_paragraph()
            last_run = last_paragraph.add_run()
            last_run.add_break(WD_BREAK.PAGE)

        # Сохраняем обработанный документ в новый буфер
        processed_buffer = io.BytesIO()
        temp_doc.save(processed_buffer)
        processed_buffer.seek(0)

        processed_docs.append({
            'buffer': processed_buffer,
            'item': doc_data['item'],
            'nameplate_number': doc_data['nameplate_number']
        })

    # Создаем основной документ из первого обработанного
    main_doc = Document(processed_docs[0]['buffer'])
    composer = Composer(main_doc)

    # Добавляем остальные обработанные документы
    for doc_data in processed_docs[1:]:
        temp_doc = Document(doc_data['buffer'])
        composer.append(temp_doc)

    # Сохраняем результат
    final_buffer = io.BytesIO()
    composer.save(final_buffer)
    final_buffer.seek(0)

    return final_buffer